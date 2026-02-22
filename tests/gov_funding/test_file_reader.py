"""file_reader.py 단위 테스트 — PDF/DOCX 텍스트 추출"""
import io

import pytest

from src.gov_funding.utils.file_reader import (
    extract_text_from_docx,
    extract_text_from_file,
    extract_text_from_pdf,
)


# --- PDF 추출 테스트 ---


def _make_pdf_bytes(text: str) -> bytes:
    """테스트용 PDF 바이트 생성 (pdfplumber로 읽을 수 있는 형식)"""
    from pypdfium2 import PdfDocument

    pdf = PdfDocument.new()
    page = pdf.new_page(595, 842)  # A4 크기

    # pypdfium2는 텍스트 삽입이 복잡하므로 fpdf2 대신 reportlab 패턴 사용
    # 간단한 방법: pdfplumber가 읽을 수 있는 최소 PDF 직접 생성
    buf = io.BytesIO()
    pdf.save(buf)
    pdf.close()
    return buf.getvalue()


def _make_simple_pdf(text: str) -> bytes:
    """최소한의 PDF 바이트 직접 생성 (텍스트 포함)"""
    # PDF 1.4 최소 구조 — 텍스트 스트림 포함
    stream_content = f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET"
    stream_bytes = stream_content.encode("latin-1")
    stream_len = len(stream_bytes)

    pdf = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length {stream_len} >>
stream
{stream_content}
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
{str(334 + stream_len).zfill(10)} 00000 n

trailer
<< /Size 6 /Root 1 0 R >>
startxref
{399 + stream_len}
%%EOF"""
    return pdf.encode("latin-1")


def test_extract_text_from_pdf_basic():
    """PDF에서 텍스트 추출 기본 테스트"""
    pdf_bytes = _make_simple_pdf("Hello PDF World")
    result = extract_text_from_pdf(pdf_bytes)
    assert "Hello PDF World" in result


def test_extract_text_from_pdf_empty():
    """빈 PDF에서 빈 텍스트 반환"""
    # 텍스트 없는 최소 PDF
    pdf_bytes = _make_simple_pdf("")
    result = extract_text_from_pdf(pdf_bytes)
    assert isinstance(result, str)


def test_extract_text_from_pdf_invalid_bytes():
    """잘못된 바이트에서 빈 문자열 반환"""
    result = extract_text_from_pdf(b"not a pdf file")
    assert result == ""


# --- DOCX 추출 테스트 ---


def _make_docx_bytes(paragraphs: list[str], table_data: list[list[str]] | None = None) -> bytes:
    """테스트용 DOCX 바이트 생성"""
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)

    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_text_from_docx_paragraphs():
    """DOCX 단락 텍스트 추출"""
    docx_bytes = _make_docx_bytes(["첫 번째 단락", "두 번째 단락"])
    result = extract_text_from_docx(docx_bytes)
    assert "첫 번째 단락" in result
    assert "두 번째 단락" in result


def test_extract_text_from_docx_with_table():
    """DOCX 테이블 텍스트 추출"""
    docx_bytes = _make_docx_bytes(
        ["본문 텍스트"],
        table_data=[["항목", "값"], ["이름", "테스트"]],
    )
    result = extract_text_from_docx(docx_bytes)
    assert "본문 텍스트" in result
    assert "항목" in result
    assert "테스트" in result


def test_extract_text_from_docx_empty():
    """빈 DOCX에서 빈 텍스트 반환"""
    docx_bytes = _make_docx_bytes([])
    result = extract_text_from_docx(docx_bytes)
    assert result == ""


def test_extract_text_from_docx_invalid_bytes():
    """잘못된 바이트에서 빈 문자열 반환"""
    result = extract_text_from_docx(b"not a docx file")
    assert result == ""


# --- extract_text_from_file 디스패처 테스트 ---


def test_dispatch_pdf():
    """확장자 .pdf로 PDF 추출 호출"""
    pdf_bytes = _make_simple_pdf("dispatch test")
    result = extract_text_from_file(pdf_bytes, "문서.pdf")
    assert "dispatch test" in result


def test_dispatch_docx():
    """확장자 .docx로 DOCX 추출 호출"""
    docx_bytes = _make_docx_bytes(["디스패치 테스트"])
    result = extract_text_from_file(docx_bytes, "문서.docx")
    assert "디스패치 테스트" in result


def test_dispatch_unsupported():
    """지원하지 않는 확장자는 빈 문자열 반환"""
    result = extract_text_from_file(b"data", "이미지.png")
    assert result == ""


def test_dispatch_case_insensitive():
    """확장자 대소문자 무관"""
    pdf_bytes = _make_simple_pdf("case test")
    result = extract_text_from_file(pdf_bytes, "문서.PDF")
    assert "case test" in result
