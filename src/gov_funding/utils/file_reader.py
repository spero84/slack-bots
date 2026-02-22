"""파일 텍스트 추출 유틸리티

지원 형식:
- HWP (바이너리): olefile로 OLE2 컨테이너 → BodyText/Section 스트림에서 zlib 압축 해제 → 텍스트 추출
- HWPX (XML ZIP): zipfile로 압축 해제 → Contents/section*.xml에서 lxml로 텍스트 노드 추출
- PDF: pdfplumber로 레이아웃 보존 텍스트 추출
- DOCX: python-docx로 단락 및 테이블 텍스트 추출
"""
import io
import re
import struct
import zlib
import zipfile
from typing import Optional

from lxml import etree

from ..utils.logger import logger

try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    logger.warning("olefile not available - HWP binary parsing disabled")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not available - PDF parsing disabled")

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX parsing disabled")


def extract_text_from_hwp(file_bytes: bytes) -> str:
    """HWP 바이너리 파일에서 텍스트 추출

    HWP 5.0 포맷: OLE2 컨테이너 내 BodyText/Section0~N 스트림에서
    zlib 압축된 레코드를 파싱하여 텍스트 추출.

    Args:
        file_bytes: HWP 파일 바이트

    Returns:
        추출된 텍스트
    """
    if not OLEFILE_AVAILABLE:
        logger.error("olefile not installed - cannot parse HWP")
        return ""

    try:
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"HWP OLE 파일 열기 실패: {e}")
        return ""

    text_parts = []

    try:
        # FileHeader에서 압축 여부 확인
        is_compressed = True
        if ole.exists("FileHeader"):
            header = ole.openstream("FileHeader").read()
            if len(header) >= 40:
                # 바이트 36의 비트 0이 압축 플래그
                flags = struct.unpack_from("<I", header, 36)[0]
                is_compressed = bool(flags & 0x01)

        # BodyText/Section0, Section1, ... 순서대로 처리
        section_idx = 0
        while True:
            section_name = f"BodyText/Section{section_idx}"
            if not ole.exists(section_name):
                break

            raw_data = ole.openstream(section_name).read()

            if is_compressed:
                try:
                    raw_data = zlib.decompress(raw_data, -15)
                except zlib.error:
                    try:
                        raw_data = zlib.decompress(raw_data)
                    except zlib.error:
                        logger.debug(f"HWP Section{section_idx} zlib 압축 해제 실패")
                        section_idx += 1
                        continue

            # 레코드 파싱하여 텍스트 추출
            section_text = _parse_hwp_records(raw_data)
            if section_text:
                text_parts.append(section_text)

            section_idx += 1

    except Exception as e:
        logger.error(f"HWP 텍스트 추출 오류: {e}")
    finally:
        ole.close()

    return "\n".join(text_parts)


def _parse_hwp_records(data: bytes) -> str:
    """HWP 바이너리 레코드에서 텍스트 추출

    HWP 레코드 헤더: 4바이트 (태그ID 10비트 + 레벨 10비트 + 크기 12비트)
    태그 ID 67 = HWPTAG_PARA_TEXT (텍스트 레코드)
    """
    text_parts = []
    offset = 0

    while offset < len(data) - 4:
        try:
            header = struct.unpack_from("<I", data, offset)[0]
            tag_id = header & 0x3FF
            # level = (header >> 10) & 0x3FF
            size = (header >> 20) & 0xFFF

            offset += 4

            # 크기가 0xFFF이면 확장 크기 사용
            if size == 0xFFF:
                if offset + 4 > len(data):
                    break
                size = struct.unpack_from("<I", data, offset)[0]
                offset += 4

            if offset + size > len(data):
                break

            record_data = data[offset:offset + size]
            offset += size

            # HWPTAG_PARA_TEXT (tag_id = 67)
            if tag_id == 67 and size >= 2:
                text = _extract_text_from_para(record_data)
                if text:
                    text_parts.append(text)

        except (struct.error, IndexError):
            break

    return "\n".join(text_parts)


def _extract_text_from_para(record_data: bytes) -> str:
    """PARA_TEXT 레코드에서 텍스트 추출

    UTF-16LE 인코딩, 특수 제어 문자 처리:
    - 0~31 범위의 유니코드는 HWP 제어 문자
    - 0x00: 무효
    - 0x01~0x03: 인라인 컨트롤 (각각 추가 바이트 소비)
    - 0x04~0x09: 확장 컨트롤 (각각 추가 바이트 소비)
    - 0x0A: 줄바꿈
    - 0x0D: 단락 끝
    """
    chars = []
    i = 0
    length = len(record_data)

    while i < length - 1:
        code = struct.unpack_from("<H", record_data, i)[0]
        i += 2

        if code == 0:
            continue
        elif code in (1, 2, 3):
            # 인라인 확장 문자: 추가 바이트 스킵 (각각 다른 크기)
            i += 14  # 일반적으로 인라인 컨트롤은 추가 14바이트
            if i > length:
                break
        elif 4 <= code <= 9:
            # 확장 컨트롤: 추가 바이트 스킵
            i += 14
            if i > length:
                break
        elif code == 10:
            chars.append("\n")
        elif code == 13:
            chars.append("\n")
        elif code == 24:
            # 하이픈
            chars.append("-")
        elif code < 32:
            # 기타 제어 문자 무시
            continue
        else:
            chars.append(chr(code))

    text = "".join(chars).strip()
    # 연속 빈 줄 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def extract_text_from_hwpx(file_bytes: bytes) -> str:
    """HWPX (XML ZIP) 파일에서 텍스트 추출

    HWPX는 OOXML 기반으로 ZIP 컨테이너 내 Contents/section*.xml 파일에
    텍스트가 저장됨.

    Args:
        file_bytes: HWPX 파일 바이트

    Returns:
        추출된 텍스트
    """
    text_parts = []

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            # Contents/section0.xml, section1.xml, ... 순서대로 처리
            section_files = sorted([
                name for name in zf.namelist()
                if re.match(r'Contents/section\d+\.xml', name, re.IGNORECASE)
            ])

            if not section_files:
                # 일부 HWPX는 Contents/content.hpf 등의 다른 구조일 수 있음
                section_files = sorted([
                    name for name in zf.namelist()
                    if name.lower().endswith('.xml') and 'section' in name.lower()
                ])

            for section_file in section_files:
                try:
                    xml_data = zf.read(section_file)
                    section_text = _parse_hwpx_section(xml_data)
                    if section_text:
                        text_parts.append(section_text)
                except Exception as e:
                    logger.debug(f"HWPX {section_file} 파싱 오류: {e}")

    except zipfile.BadZipFile:
        logger.error("HWPX 파일이 유효한 ZIP이 아님")
    except Exception as e:
        logger.error(f"HWPX 텍스트 추출 오류: {e}")

    return "\n".join(text_parts)


def _parse_hwpx_section(xml_data: bytes) -> str:
    """HWPX section XML에서 텍스트 추출

    네임스페이스: http://www.hancom.co.kr/hwpml/2011/paragraph
    텍스트 노드: hp:t (또는 {ns}t)
    """
    text_parts = []

    try:
        root = etree.fromstring(xml_data)

        # 네임스페이스 자동 감지
        nsmap = root.nsmap
        # hp 또는 기본 네임스페이스에서 텍스트 추출
        # 모든 텍스트 노드를 깊이 우선으로 수집
        for elem in root.iter():
            tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ""

            # <t> 태그가 텍스트 노드
            if tag == "t" and elem.text:
                text_parts.append(elem.text)
            # <linesegarray> 등에서 줄바꿈 처리
            elif tag in ("p", "subList"):
                text_parts.append("\n")

    except etree.XMLSyntaxError as e:
        logger.debug(f"HWPX XML 파싱 오류: {e}")
    except Exception as e:
        logger.debug(f"HWPX 텍스트 추출 오류: {e}")

    text = "".join(text_parts).strip()
    # 연속 빈 줄 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """PDF에서 텍스트 추출 (pdfplumber 사용)

    Args:
        file_bytes: PDF 파일 바이트

    Returns:
        추출된 텍스트
    """
    if not PDFPLUMBER_AVAILABLE:
        logger.error("pdfplumber not installed - cannot parse PDF")
        return ""

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            parts = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error(f"PDF 텍스트 추출 오류: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """DOCX에서 텍스트 추출 (python-docx 사용)

    Args:
        file_bytes: DOCX 파일 바이트

    Returns:
        추출된 텍스트
    """
    if not PYTHON_DOCX_AVAILABLE:
        logger.error("python-docx not installed - cannot parse DOCX")
        return ""

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 테이블 텍스트도 추출
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error(f"DOCX 텍스트 추출 오류: {e}")
        return ""


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """파일 확장자에 따라 적절한 텍스트 추출 함수 호출

    Args:
        file_bytes: 파일 바이트
        filename: 파일명 (확장자 판별용)

    Returns:
        추출된 텍스트
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "hwpx":
        return extract_text_from_hwpx(file_bytes)
    elif ext == "hwp":
        return extract_text_from_hwp(file_bytes)
    elif ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        logger.warning(f"지원하지 않는 파일 형식: {filename}")
        return ""
